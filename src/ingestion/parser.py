"""Document Parser for NEPPA Ingestion Pipeline.

This module handles parsing PDF and DOCX documents, extracting text content,
and tagging them with appropriate metadata based on source folder structure.
"""

import json
import logging
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import fitz  # PyMuPDF
from docx import Document

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)


class DocumentMetadata:
    """Metadata schema for parsed documents.

    Attributes:
        source_type: Category of document (National, Local, Governance, Legal)
        organization: Source organization (NICE, CPICS, NHS England, etc.)
        file_name: Original filename
        file_path: Relative path to the source file
        clinical_area: Clinical domain (e.g., "Diabetes", "General Governance")
        last_updated: Date string if extractable (optional)
        sortable_date: Date in YYYYMMDD format for sorting (default: 20220101)
        priority_score: Retrieval priority score (Local=1.0, National=0.8, Legal/Governance=0.5)
        is_presentation: Boolean flag indicating if document is a PowerPoint presentation
    """

    def __init__(
        self,
        source_type: str,
        organization: str,
        file_name: str,
        file_path: str,
        clinical_area: str,
        last_updated: Optional[str] = None,
        sortable_date: str = "20220101",
        priority_score: float = 0.5,
        is_presentation: bool = False,
    ):
        self.source_type = source_type
        self.organization = organization
        self.file_name = file_name
        self.file_path = file_path
        self.clinical_area = clinical_area
        self.last_updated = last_updated
        self.sortable_date = sortable_date
        self.priority_score = priority_score
        self.is_presentation = is_presentation

    def to_dict(self, context_header: Optional[str] = None) -> Dict:
        """Convert metadata to dictionary for JSON serialization.
        
        Args:
            context_header: Optional section heading for this chunk.
        
        Returns:
            Dictionary with all metadata fields including context_header.
        """
        metadata_dict = {
            "source_type": self.source_type,
            "organization": self.organization,
            "file_name": self.file_name,
            "file_path": self.file_path,
            "clinical_area": self.clinical_area,
            "last_updated": self.last_updated,
            "sortable_date": self.sortable_date,
            "priority_score": self.priority_score,
            "is_presentation": self.is_presentation,
        }
        if context_header:
            metadata_dict["context_header"] = context_header
        return metadata_dict

    def __repr__(self) -> str:
        return (
            f"DocumentMetadata(source_type={self.source_type}, "
            f"organization={self.organization}, file_name={self.file_name})"
        )


class DocumentParser:
    """Parser for PDF and DOCX documents with metadata extraction.

    This parser extracts text content from PDF and DOCX files and generates
    metadata based on the folder structure and filename patterns.
    """

    # Folder mapping: folder name prefix -> source_type
    SOURCE_TYPE_MAP = {
        "01_National": "National",
        "02_Local": "Local",
        "03_Governance": "Governance",
        "04_IFR_process": "Legal",
    }

    # Organization inference from filename/folder
    ORGANIZATION_MAP = {
        "NICE": ["nice", "ng28", "type-2-diabetes"],
        "CPICS": ["cpics", "cambridgeshire", "peterborough", "les"],
        "NHS England": ["constitution", "nhs england"],
    }
    
    # Governance document keywords for enhanced detection
    GOVERNANCE_KEYWORDS = [
        "nhs england",
        "department of health",
        "integrated care board",
        "icb",
        "nhs england and nhs improvement",
        "commissioning",
    ]
    
    # PowerPoint detection keywords (must be in filename)
    PPT_KEYWORDS = [
        "powerpoint",
        "presentation",
        "ppt",
        "slides",
        "slide",  # Singular form
    ]
    
    # Medical acronyms for normalization
    MEDICAL_ACRONYMS = {
        "T2D": "type 2 diabetes",
        "T2DM": "type 2 diabetes mellitus",
        "CGM": "continuous glucose monitoring",
        "IFR": "individual funding request",
        "ICB": "integrated care board",
        "CPICS": "cambridgeshire and peterborough integrated care system",
        "SGLT2": "sodium-glucose cotransporter 2",
        "SGLT2i": "sodium-glucose cotransporter 2 inhibitor",
        "GLP-1": "glucagon-like peptide-1",
        "DPP-4": "dipeptidyl peptidase-4",
        "eGFR": "estimated glomerular filtration rate",
        "HbA1c": "glycated hemoglobin",
        "BMI": "body mass index",
        "CKD": "chronic kidney disease",
        "HF": "heart failure",
        "DKA": "diabetic ketoacidosis",
        "ACE": "angiotensin-converting enzyme",
        "ARB": "angiotensin receptor blocker",
    }
    
    # Drug names for context header enhancement
    DRUG_NAMES = [
        "dapagliflozin", "tirzepatide", "empagliflozin", "insulin", "metformin",
        "glp-1", "sglt2", "sglt2i", "dpp-4", "canagliflozin", "semaglutide"
    ]
    
    # Condition names for context header enhancement
    CONDITION_NAMES = [
        "diabetes", "type 2 diabetes", "t2d", "t2dm", "ckd", "chronic kidney disease",
        "heart failure", "hf", "diabetic ketoacidosis", "dka", "hypoglycaemia"
    ]

    def __init__(self, data_root: Path):
        """Initialize the parser with data root directory.

        Args:
            data_root: Path to the data/raw directory containing document folders.
        """
        self.data_root = Path(data_root)
        if not self.data_root.exists():
            raise ValueError(f"Data root directory does not exist: {data_root}")

    def _detect_section_headings(self, text: str) -> List[Tuple[int, str]]:
        """Detect section headings in text with improved pattern matching.
        
        Looks for patterns like:
        - Lines that are all caps or title case
        - Lines ending with colon
        - Lines that are short (typically < 100 chars)
        - Lines followed by blank lines
        - Numbered sections (e.g., "1. Introduction", "Section 2:", "2.1 ")
        - Medical document patterns (e.g., "NICE Guidance", "Contraindications")
        
        Args:
            text: Full document text.
        
        Returns:
            List of tuples (line_index, heading_text) for detected headings.
        """
        lines = text.split('\n')
        headings = []
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            if not line_stripped:
                continue
            
            # Skip very long lines (likely not headings)
            if len(line_stripped) > 120:
                continue
            
            # Skip lines that are clearly not headings (contain sentence-ending punctuation mid-line)
            if re.search(r'[.!?]\s+[a-z]', line_stripped):
                continue
            
            # Pattern 1: Lines ending with colon (common heading pattern)
            # But exclude if it's a URL or email
            if (line_stripped.endswith(':') and 
                len(line_stripped) < 100 and
                '://' not in line_stripped and
                '@' not in line_stripped):
                # Clean up the heading (remove extra spaces, normalize)
                clean_heading = re.sub(r'\s+', ' ', line_stripped)
                headings.append((i, clean_heading))
                continue
            
            # Pattern 2: Numbered sections (e.g., "1. ", "Section 2:", "2.1 ", "1.1.1")
            numbered_pattern = re.match(
                r'^(\d+\.?\s+|\d+\.\d+\.?\s+|\d+\.\d+\.\d+\.?\s+|[A-Z][a-z]+\s+\d+[.:]\s+)',
                line_stripped
            )
            if numbered_pattern:
                if len(line_stripped) < 120:
                    clean_heading = re.sub(r'\s+', ' ', line_stripped)
                    headings.append((i, clean_heading))
                continue
            
            # Pattern 3: All caps short lines (likely headings)
            if line_stripped.isupper() and len(line_stripped) < 80 and len(line_stripped) > 3:
                # Check if next line is blank or starts with lowercase (indicates heading)
                if i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    if not next_line or (next_line and next_line[0].islower()):
                        clean_heading = re.sub(r'\s+', ' ', line_stripped)
                        headings.append((i, clean_heading))
                continue
            
            # Pattern 4: Title case lines that are short and followed by content
            if (line_stripped.istitle() and 
                len(line_stripped) < 100 and 
                len(line_stripped) > 3 and
                not line_stripped.endswith('.') and
                not line_stripped.endswith(',') and
                not line_stripped.endswith(';')):
                # Check if next line is blank or starts with lowercase
                if i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    if not next_line or (next_line and next_line[0].islower()):
                        # Additional check: line should not be a full sentence
                        if not re.search(r'[.!?]$', line_stripped):
                            clean_heading = re.sub(r'\s+', ' ', line_stripped)
                            headings.append((i, clean_heading))
                continue
            
            # Pattern 5: Common heading keywords (medical/policy documents)
            heading_keywords = [
                'what is', 'where', 'when', 'who', 'how', 'why',
                'introduction', 'overview', 'summary', 'conclusion',
                'background', 'method', 'result', 'discussion',
                'recommendation', 'guidance', 'policy', 'procedure',
                'contraindication', 'indication', 'dosage', 'administration',
                'adverse effect', 'side effect', 'monitoring', 'preparation',
                'pregnancy', 'breastfeeding', 'interaction', 'cautions',
                'licenced indication', 'prescribing', 'stopping therapy',
                'advice and support', 'reference', 'document ratification',
                'nice guidance', 'nice technology', 'drug interactions',
                'preparations and dosage', 'contraindications and cautions'
            ]
            line_lower = line_stripped.lower()
            if any(keyword in line_lower for keyword in heading_keywords):
                # Check if next line is blank or starts new content (indicates heading)
                is_heading = False
                if i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    if not next_line or (next_line and next_line[0].islower()):
                        is_heading = True
                else:
                    is_heading = True  # Last line could be a heading
                
                if is_heading and len(line_stripped) < 120:
                    clean_heading = re.sub(r'\s+', ' ', line_stripped)
                    headings.append((i, clean_heading))
                    continue
            
            # Pattern 6: Lines containing drug names (for better context)
            line_lower = line_stripped.lower()
            if any(drug in line_lower for drug in self.DRUG_NAMES):
                # If line is short and looks like a heading
                if (len(line_stripped) < 100 and 
                    (line_stripped.endswith(':') or line_stripped.istitle() or 
                     (i + 1 < len(lines) and not lines[i + 1].strip()))):
                    clean_heading = re.sub(r'\s+', ' ', line_stripped)
                    headings.append((i, clean_heading))
                    continue
        
        # Remove duplicate consecutive headings
        unique_headings = []
        last_heading = None
        for heading in headings:
            if heading[1] != last_heading:
                unique_headings.append(heading)
                last_heading = heading[1]
        
        # Filter out weak headings (partial lines, measurements, etc.)
        filtered_headings = []
        weak_patterns = [
            r'^\d+\s*(ml|mg|units?|mm|cm|kg|g|%)',  # Measurements
            r'^\d+\.\d+',  # Decimal numbers
            r'^[a-z]+\s+[a-z]+\s+[a-z]+$',  # Very short generic phrases
            r'^(and|or|the|a|an)\s+',  # Starting with articles
            r'^page\s+\d+',  # Page numbers
        ]
        
        for heading in unique_headings:
            heading_text = heading[1]
            is_weak = any(re.match(pattern, heading_text.lower()) for pattern in weak_patterns)
            
            # Also filter if it's too short and doesn't contain meaningful words
            if len(heading_text) < 10 and not any(
                keyword in heading_text.lower() 
                for keyword in ['guidance', 'dosage', 'drug', 'contraindication', 'interaction']
            ):
                is_weak = True
            
            if not is_weak:
                filtered_headings.append(heading)
            else:
                # Try to find a better heading nearby (look back up to 3 lines)
                line_idx = heading[0]
                for lookback in range(1, min(4, line_idx + 1)):
                    prev_line_idx = line_idx - lookback
                    if prev_line_idx >= 0:
                        prev_line = lines[prev_line_idx].strip()
                        if prev_line and len(prev_line) < 100:
                            # Check if previous line is a better heading
                            prev_lower = prev_line.lower()
                            if any(keyword in prev_lower for keyword in 
                                   ['guidance', 'dosage', 'drug', 'contraindication', 'interaction', 
                                    'nice', 'preparation', 'administration']):
                                if prev_line not in [h[1] for h in filtered_headings]:
                                    filtered_headings.append((prev_line_idx, prev_line))
                                    break
        
        return filtered_headings

    def _normalize_acronyms(self, text: str) -> str:
        """Normalize medical acronyms by expanding them in the text.
        
        Adds full forms after first occurrence of acronyms for better search recall.
        
        Args:
            text: Text content to normalize.
        
        Returns:
            Text with acronym expansions added.
        """
        normalized_text = text
        
        # Sort acronyms by length (longest first) to avoid partial matches
        sorted_acronyms = sorted(self.MEDICAL_ACRONYMS.items(), key=lambda x: len(x[0]), reverse=True)
        
        for acronym, full_form in sorted_acronyms:
            # Pattern to match acronym as whole word (case-insensitive)
            pattern = rf'\b{re.escape(acronym)}\b'
            
            # Check if acronym exists in text
            if re.search(pattern, normalized_text, re.IGNORECASE):
                # Check if full form already exists nearby (within 50 chars)
                matches = list(re.finditer(pattern, normalized_text, re.IGNORECASE))
                for match in matches:
                    start_pos = match.start()
                    end_pos = match.end()
                    
                    # Check if full form exists within 50 characters before or after
                    context_start = max(0, start_pos - 50)
                    context_end = min(len(normalized_text), end_pos + 50)
                    context = normalized_text[context_start:context_end].lower()
                    
                    if full_form.lower() not in context:
                        # Add full form in parentheses after first occurrence
                        normalized_text = (
                            normalized_text[:end_pos] + 
                            f" ({full_form})" + 
                            normalized_text[end_pos:]
                        )
                        # Only expand first occurrence per acronym to avoid clutter
                        break
        
        return normalized_text

    def _chunk_presentation_pages(
        self,
        page_texts: List[Tuple[int, str]],
        metadata: DocumentMetadata
    ) -> List[Dict]:
        """Chunk a PowerPoint presentation using hard page breaks.
        
        Each page becomes its own chunk with the slide title prepended.
        
        Args:
            page_texts: List of (page_num, page_text) tuples.
            metadata: DocumentMetadata object.
        
        Returns:
            List of chunk dictionaries with text, metadata, and chunk_id.
        """
        chunks = []
        
        for page_num, page_text in page_texts:
            if not page_text.strip():
                continue
            
            # Extract slide title (first non-empty line of the page)
            lines = page_text.split('\n')
            slide_title = None
            for line in lines:
                line_stripped = line.strip()
                if line_stripped and len(line_stripped) > 3:
                    # Clean up the title
                    slide_title = re.sub(r'\s+', ' ', line_stripped)
                    if len(slide_title) > 100:
                        slide_title = slide_title[:100] + "..."
                    break
            
            # Prepend slide title to page text
            if slide_title:
                chunk_text = f"Slide: {slide_title}\n\n{page_text.strip()}"
                context_header = slide_title
            else:
                chunk_text = page_text.strip()
                context_header = f"Slide {page_num + 1}"
            
            # Normalize acronyms in chunk text
            chunk_text = self._normalize_acronyms(chunk_text)
            
            chunk_num = len(chunks) + 1
            chunk_id = f"{metadata.source_type}_{Path(metadata.file_name).stem}_slide{page_num + 1}"
            
            chunk_metadata = metadata.to_dict(context_header=context_header)
            
            chunks.append({
                "text": chunk_text,
                "metadata": chunk_metadata,
                "chunk_id": chunk_id,
            })
        
        return chunks

    def _chunk_with_context(
        self, 
        text: str, 
        metadata: DocumentMetadata,
        chunk_size: int = 1000,
        overlap: int = 200
    ) -> List[Dict]:
        """Split text into chunks with context headers.
        
        Args:
            text: Full document text.
            metadata: DocumentMetadata object.
            chunk_size: Target chunk size in characters.
            overlap: Overlap between chunks in characters.
        
        Returns:
            List of chunk dictionaries with text, metadata, and chunk_id.
        """
        # Detect section headings
        headings = self._detect_section_headings(text)
        lines = text.split('\n')
        
        chunks = []
        current_chunk = []
        current_chunk_size = 0
        current_heading = None
        heading_index = 0
        
        for i, line in enumerate(lines):
            # Check if this line is a heading
            if heading_index < len(headings) and headings[heading_index][0] == i:
                current_heading = headings[heading_index][1]
                heading_index += 1
            
            # Add line to current chunk
            line_with_newline = line + '\n'
            current_chunk.append(line_with_newline)
            current_chunk_size += len(line_with_newline)
            
            # Create chunk if we've reached the target size
            if current_chunk_size >= chunk_size:
                chunk_text = ''.join(current_chunk).strip()
                if chunk_text:
                    # Normalize acronyms in chunk text
                    chunk_text = self._normalize_acronyms(chunk_text)
                    
                    chunk_num = len(chunks) + 1
                    chunk_id = f"{metadata.source_type}_{Path(metadata.file_name).stem}_chunk{chunk_num}"
                    
                    chunk_metadata = metadata.to_dict(context_header=current_heading)
                    
                    chunks.append({
                        "text": chunk_text,
                        "metadata": chunk_metadata,
                        "chunk_id": chunk_id,
                    })
                
                # Start new chunk with overlap
                if overlap > 0 and current_chunk:
                    # Keep last few lines for overlap
                    overlap_text = ''.join(current_chunk[-min(5, len(current_chunk)):])
                    current_chunk = [overlap_text] if len(overlap_text) < overlap else []
                    current_chunk_size = len(overlap_text)
                else:
                    current_chunk = []
                    current_chunk_size = 0
        
        # Add final chunk if there's remaining text
        if current_chunk:
            chunk_text = ''.join(current_chunk).strip()
            if chunk_text:
                # Normalize acronyms in chunk text
                chunk_text = self._normalize_acronyms(chunk_text)
                
                chunk_num = len(chunks) + 1
                chunk_id = f"{metadata.source_type}_{Path(metadata.file_name).stem}_chunk{chunk_num}"
                
                chunk_metadata = metadata.to_dict(context_header=current_heading)
                
                chunks.append({
                    "text": chunk_text,
                    "metadata": chunk_metadata,
                    "chunk_id": chunk_id,
                })
        
        # If no chunks were created (very short document), create one chunk
        if not chunks and text.strip():
            chunk_text = self._normalize_acronyms(text.strip())
            chunk_id = f"{metadata.source_type}_{Path(metadata.file_name).stem}_chunk1"
            chunk_metadata = metadata.to_dict(context_header=None)
            chunks.append({
                "text": chunk_text,
                "metadata": chunk_metadata,
                "chunk_id": chunk_id,
            })
        
        return chunks

    def _detect_presentation(self, file_path: Path) -> bool:
        """Detect if a PDF is likely a PowerPoint presentation export.
        
        Args:
            file_path: Path to the document file.
        
        Returns:
            True if document appears to be a PowerPoint presentation.
        """
        file_lower = file_path.name.lower()
        
        # Check filename for PowerPoint keywords
        if any(keyword in file_lower for keyword in self.PPT_KEYWORDS):
            return True
        
        return False
    
    def _enhance_governance_organization(self, text: str, current_org: str) -> str:
        """Enhance organization detection for Governance documents by scanning text.
        
        Args:
            text: Document text content (first 2000 characters will be scanned).
            current_org: Currently detected organization.
        
        Returns:
            Enhanced organization name or current_org if no match found.
        """
        if current_org != "Unknown":
            return current_org
        
        # Scan first 2000 characters for governance keywords
        scan_text = text[:2000].lower()
        
        for keyword in self.GOVERNANCE_KEYWORDS:
            if keyword in scan_text:
                # Determine specific organization
                if "nhs england" in scan_text or "nhs england and nhs improvement" in scan_text:
                    return "NHS England"
                elif "department of health" in scan_text:
                    return "Department of Health"
                elif "integrated care board" in scan_text or "icb" in scan_text:
                    # Could be CPICS or another ICB, but default to NHS England for governance
                    return "NHS England"
                elif "commissioning" in scan_text:
                    return "NHS England"
        
        return current_org
    
    def _extract_sortable_date(self, last_updated: Optional[str], file_path: Path) -> str:
        """Extract sortable date in YYYYMMDD format.
        
        Args:
            last_updated: Date string in format YYYY-MM or None.
            file_path: Path to extract date from filename if needed.
        
        Returns:
            Date string in YYYYMMDD format, default "20220101" if not found.
        """
        # If we have last_updated in YYYY-MM format, convert to YYYYMMDD
        if last_updated and len(last_updated) == 7 and last_updated[4] == '-':
            year = last_updated[:4]
            month = last_updated[5:7]
            return f"{year}{month}01"  # Use first day of month
        
        # Try to extract from filename
        stem = file_path.stem
        
        # Try DDMMYYYY format (e.g., "27062024-")
        if len(stem) >= 8 and stem[:8].isdigit():
            day = stem[:2]
            month = stem[2:4]
            year = stem[4:8]
            if 2000 <= int(year) <= 2100 and 1 <= int(month) <= 12 and 1 <= int(day) <= 31:
                return f"{year}{month}{day}"
        
        # Try YYYYMM format (e.g., "202310-")
        if len(stem) >= 6 and stem[:6].isdigit():
            year = stem[:4]
            month = stem[4:6]
            if 2000 <= int(year) <= 2100 and 1 <= int(month) <= 12:
                return f"{year}{month}01"
        
        return "20220101"  # Default date
    
    def _infer_metadata_from_path(self, file_path: Path) -> DocumentMetadata:
        """Infer metadata from file path and name.

        Args:
            file_path: Full path to the document file.

        Returns:
            DocumentMetadata object with inferred attributes.
        """
        # Extract relative path from data_root
        try:
            relative_path = file_path.relative_to(self.data_root)
        except ValueError:
            relative_path = Path(file_path.name)

        # Get parent folder name (e.g., "01_National", "02_Local")
        parent_folder = relative_path.parent.name
        source_type = self.SOURCE_TYPE_MAP.get(parent_folder, "Unknown")

        # Infer organization from filename and folder structure
        file_lower = file_path.name.lower()
        organization = "Unknown"
        
        # Default organization based on source type
        if source_type == "Local":
            organization = "CPICS"  # Local documents are from Cambridgeshire & Peterborough ICB
        elif source_type == "Legal":
            organization = "NHS England"
        elif source_type == "Governance":
            # Default to NHS England for Governance, but will be enhanced by text scanning
            organization = "NHS England"
        
        # Override with keyword matching if found
        for org, keywords in self.ORGANIZATION_MAP.items():
            if any(keyword in file_lower for keyword in keywords):
                organization = org
                break

        # Infer clinical area (default to "General" if not diabetes-related)
        clinical_area = "General Governance"
        if any(
            keyword in file_lower
            for keyword in ["diabetes", "glucose", "insulin", "tirzepatide", "dapagliflozin"]
        ):
            clinical_area = "Diabetes"
        elif "constitution" in file_lower:
            clinical_area = "Patient Rights"
        elif "ifr" in file_lower or "funding" in file_lower:
            clinical_area = "Funding Policy"

        # Try to extract date from filename
        # Formats: "202310-" (YYYYMM), "27062024-" (DDMMYYYY)
        last_updated = None
        sortable_date = "20220101"  # Default date
        stem = file_path.stem
        
        # Try YYYYMM format first (e.g., "202310-", "202407-")
        if len(stem) >= 6 and stem[:6].isdigit():
            year = stem[:4]
            month = stem[4:6]
            # Validate: year should be reasonable (2000-2100), month 01-12
            if 2000 <= int(year) <= 2100 and 1 <= int(month) <= 12:
                last_updated = f"{year}-{month}"
                sortable_date = f"{year}{month}01"  # Use first day of month
        
        # Try DDMMYYYY format if YYYYMM didn't match (e.g., "27062024-")
        if not last_updated and len(stem) >= 8 and stem[:8].isdigit():
            day = stem[:2]
            month = stem[2:4]
            year = stem[4:8]
            # Validate: reasonable date
            if 2000 <= int(year) <= 2100 and 1 <= int(month) <= 12 and 1 <= int(day) <= 31:
                last_updated = f"{year}-{month}"
                sortable_date = f"{year}{month}{day}"
        
        # Calculate priority score based on source type
        priority_score = 0.5  # Default for Legal/Governance
        if source_type == "Local":
            priority_score = 1.0
        elif source_type == "National":
            priority_score = 0.8
        elif source_type in ["Legal", "Governance"]:
            priority_score = 0.5
        
        # Detect if this is likely a PowerPoint presentation
        is_presentation = self._detect_presentation(file_path)

        return DocumentMetadata(
            source_type=source_type,
            organization=organization,
            file_name=file_path.name,
            file_path=str(relative_path),
            clinical_area=clinical_area,
            last_updated=last_updated,
            sortable_date=sortable_date,
            priority_score=priority_score,
            is_presentation=is_presentation,
        )

    def parse_pdf(self, file_path: Path) -> Tuple[str, DocumentMetadata, Optional[List[Tuple[int, str]]]]:
        """Extract text content from a PDF file.

        Args:
            file_path: Path to the PDF file.

        Returns:
            Tuple of (extracted_text, metadata, page_texts)
            - extracted_text: Full document text
            - metadata: DocumentMetadata object
            - page_texts: List of (page_num, page_text) tuples if presentation, None otherwise

        Raises:
            ValueError: If file cannot be parsed.
        """
        try:
            doc = fitz.open(file_path)
            text_parts = []
            page_texts = []
            
            # Check if this is a presentation (PowerPoint export)
            is_presentation = self._detect_presentation(file_path)
            
            # Also check page orientation for landscape, but only if:
            # 1. Filename suggests it's a presentation, OR
            # 2. Document has many pages (presentations are typically longer)
            if not is_presentation and len(doc) > 0:
                first_page = doc[0]
                page_rect = first_page.rect
                # Landscape if width > height
                is_landscape = page_rect.width > page_rect.height
                
                # Only treat as presentation if landscape AND:
                # - Has filename keywords (already checked), OR
                # - Has many pages (>= 10) suggesting it's a slide deck
                if is_landscape and len(doc) >= 10:
                    is_presentation = True
                    logger.info(f"Detected landscape orientation with {len(doc)} pages - treating as presentation: {file_path.name}")

            for page_num, page in enumerate(doc):
                text = page.get_text()
                if text.strip():
                    text_parts.append(text)
                    if is_presentation:
                        page_texts.append((page_num, text))

            doc.close()
            full_text = "\n\n".join(text_parts)

            if not full_text.strip():
                logger.warning(f"PDF {file_path.name} appears to be empty or unreadable")
                full_text = ""

            metadata = self._infer_metadata_from_path(file_path)
            
            # Enhance organization detection for Governance documents
            if metadata.source_type == "Governance" and metadata.organization == "Unknown":
                enhanced_org = self._enhance_governance_organization(full_text, metadata.organization)
                if enhanced_org != metadata.organization:
                    logger.info(f"Enhanced organization detection for {file_path.name}: {metadata.organization} -> {enhanced_org}")
                    metadata.organization = enhanced_org
            
            # Update sortable_date if we have last_updated
            if metadata.last_updated:
                metadata.sortable_date = self._extract_sortable_date(metadata.last_updated, file_path)
            
            # Set is_presentation flag
            metadata.is_presentation = is_presentation
            
            # Return page_texts only if it's a presentation
            return full_text, metadata, page_texts if is_presentation else None

        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            raise ValueError(f"Failed to parse PDF: {file_path}") from e

    def parse_docx(self, file_path: Path) -> Tuple[str, DocumentMetadata, None]:
        """Extract text content from a DOCX file.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Tuple of (extracted_text, metadata, None)
            - extracted_text: Full document text
            - metadata: DocumentMetadata object
            - None: DOCX files don't have page-based structure

        Raises:
            ValueError: If file cannot be parsed.
        """
        try:
            doc = Document(file_path)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            full_text = "\n\n".join(text_parts)

            if not full_text.strip():
                logger.warning(f"DOCX {file_path.name} appears to be empty")
                full_text = ""

            metadata = self._infer_metadata_from_path(file_path)
            
            # Enhance organization detection for Governance documents
            if metadata.source_type == "Governance" and metadata.organization == "Unknown":
                enhanced_org = self._enhance_governance_organization(full_text, metadata.organization)
                if enhanced_org != metadata.organization:
                    logger.info(f"Enhanced organization detection for {file_path.name}: {metadata.organization} -> {enhanced_org}")
                    metadata.organization = enhanced_org
            
            # Update sortable_date if we have last_updated
            if metadata.last_updated:
                metadata.sortable_date = self._extract_sortable_date(metadata.last_updated, file_path)
            
            return full_text, metadata, None

        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise ValueError(f"Failed to parse DOCX: {file_path}") from e

    def discover_documents(self) -> List[Path]:
        """Discover all PDF and DOCX files in the data root.

        Returns:
            List of Path objects for discovered documents.
        """
        documents = []

        # Search for PDFs
        pdf_files = list(self.data_root.rglob("*.pdf"))
        documents.extend(pdf_files)

        # Search for DOCX files
        docx_files = list(self.data_root.rglob("*.docx"))
        documents.extend(docx_files)

        return sorted(documents)

    def parse_all(
        self, 
        output_dir: Optional[Path] = None,
        chunk_size: int = 1000,
        overlap: int = 200
    ) -> List[Dict]:
        """Parse all discovered documents and create chunks with context headers.

        Args:
            output_dir: Optional directory to save parsed chunks as JSON files.
                       If None, results are only returned.
            chunk_size: Target chunk size in characters (default: 1000).
            overlap: Overlap between chunks in characters (default: 200).

        Returns:
            List of dictionaries, each containing 'text', 'metadata', and 'chunk_id'.
        """
        documents = self.discover_documents()
        logger.info(f"Discovered {len(documents)} documents to parse")

        all_chunks = []

        for doc_path in documents:
            try:
                logger.info(f"Parsing: {doc_path.name}")

                if doc_path.suffix.lower() == ".pdf":
                    text, metadata, page_texts = self.parse_pdf(doc_path)
                elif doc_path.suffix.lower() == ".docx":
                    text, metadata, page_texts = self.parse_docx(doc_path)
                else:
                    logger.warning(f"Unsupported file type: {doc_path.suffix}")
                    continue

                # Create chunks - use page-based chunking for presentations
                if metadata.is_presentation and page_texts:
                    chunks = self._chunk_presentation_pages(page_texts, metadata)
                    logger.info(f"Created {len(chunks)} slide-based chunks for presentation: {doc_path.name}")
                else:
                    chunks = self._chunk_with_context(text, metadata, chunk_size, overlap)
                    logger.info(f"Created {len(chunks)} chunks for {doc_path.name}")

                all_chunks.extend(chunks)

                # Optionally save chunks to JSON file(s)
                if output_dir:
                    output_dir.mkdir(parents=True, exist_ok=True)
                    
                    # Save as single file with array of chunks
                    output_file = output_dir / f"{doc_path.stem}_chunks.json"
                    with open(output_file, "w", encoding="utf-8") as f:
                        json.dump(chunks, f, indent=2, ensure_ascii=False)
                    logger.info(f"Saved {len(chunks)} chunks to: {output_file}")

            except Exception as e:
                logger.error(f"Failed to parse {doc_path}: {e}")
                continue

        logger.info(f"Successfully parsed {len(documents)} documents, created {len(all_chunks)} total chunks")
        return all_chunks

